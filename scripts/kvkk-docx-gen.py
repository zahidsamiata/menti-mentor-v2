#!/usr/bin/env python3
"""KVKK belge paketi markdown -> tek profesyonel .docx (avukat sunumu).
Markdown CANONICAL kaynaktır; bu script türev .docx üretir. Yeniden üretmek için:
  python scripts/kvkk-docx-gen.py
Kaynak klasör: docs/kararlar/konu/kvkk-metinleri/  Çıktı: aynı klasör/KVKK-BELGE-PAKETI-2026-08-25.docx
"""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.join(os.path.dirname(__file__), "..", "docs", "kararlar", "konu", "kvkk-metinleri")
OUT = os.path.join(BASE, "KVKK-BELGE-PAKETI-2026-08-25.docx")
ORDER = ["00-AVUKAT-KONTROL-DOSYASI.md","01-aydinlatma-metni.md","02-acik-riza-metni.md",
         "03-gizlilik-politikasi.md","04-cerez-politikasi.md","05-saklama-imha-politikasi.md",
         "06-ilgili-kisi-basvuru-formu.md","07-kullanim-kosullari.md","08-veri-isleyen-sozlesmesi-sablonu.md"]
TITLES = {"00-AVUKAT-KONTROL-DOSYASI.md":"Avukat Kontrol Dosyası (Kapak)",
          "01-aydinlatma-metni.md":"1) Aydınlatma Metni","02-acik-riza-metni.md":"2) Açık Rıza Metni",
          "03-gizlilik-politikasi.md":"3) Gizlilik Politikası","04-cerez-politikasi.md":"4) Çerez Politikası",
          "05-saklama-imha-politikasi.md":"5) Saklama ve İmha Politikası","06-ilgili-kisi-basvuru-formu.md":"6) İlgili Kişi Başvuru Formu",
          "07-kullanim-kosullari.md":"7) Kullanım Koşulları","08-veri-isleyen-sozlesmesi-sablonu.md":"8) Veri İşleyen Sözleşmesi (Şablon)"}

# Emoji -> düz metin (Word'de bozuk glyph olmasın)
EMO = {"⚠️":"UYARI:","⚠":"UYARI:","🔴":"[KRİTİK]","⭐":"*","✅":"EVET","❌":"HAYIR","🔀":"[PR]","🟡":"[ORTA]","🟢":"","🔵":"","❓":"[SORU]",
       "📸":"","🔄":"","📋":"","🍪":"","⏳":"[BEKLİYOR]","⏭️":"","→":"->","☐":"[ ]"}
_EMORANGE=re.compile("[\U0001F000-\U0001FAFF☀-➿⬀-⯿︀-️←-⇿]")
def deemoji(s):
    for k,v in EMO.items(): s=s.replace(k,v)
    s=_EMORANGE.sub("", s)  # kalan tüm emoji/simge codepoint'lerini temizle (garanti)
    return s

def shade(cell_or_par, hexc):
    el = OxmlElement('w:shd'); el.set(qn('w:val'),'clear'); el.set(qn('w:fill'),hexc)
    cell_or_par.get_or_add_pPr().append(el) if hasattr(cell_or_par,'get_or_add_pPr') else None

def set_cell_bg(cell,hexc):
    tcpr=cell._tc.get_or_add_tcPr(); el=OxmlElement('w:shd'); el.set(qn('w:val'),'clear'); el.set(qn('w:fill'),hexc); tcpr.append(el)

MARK_H = re.compile(r'(\[HUKUKÇU KARARI[^\]]*\])')
MARK_P = re.compile(r'(\[(?:PO|KURUM) DOLDURACAK[^\]]*\])')

def add_runs(par, text):
    """inline: **bold**, `code`, [HUKUKÇU KARARI]=sarı, [PO DOLDURACAK]=gri."""
    text = deemoji(text)
    # önce marker'ları böl
    tokens = re.split(r'(\[HUKUKÇU KARARI[^\]]*\]|\[(?:PO|KURUM) DOLDURACAK[^\]]*\])', text)
    for tok in tokens:
        if not tok: continue
        if MARK_H.fullmatch(tok):
            r=par.add_run(tok); r.bold=True; r.font.highlight_color=WD_COLOR_INDEX.YELLOW; continue
        if MARK_P.fullmatch(tok):
            r=par.add_run(tok); r.italic=True; r.font.highlight_color=WD_COLOR_INDEX.GRAY_25; continue
        # bold/code parçala
        for part in re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', tok):
            if not part: continue
            if part.startswith('**') and part.endswith('**'):
                r=par.add_run(part[2:-2]); r.bold=True
            elif part.startswith('`') and part.endswith('`'):
                r=par.add_run(part[1:-1]); r.font.name='Consolas'
            else:
                par.add_run(part)

def add_field(par, instr):
    r=par.add_run(); fc=OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'),'begin'); r._r.append(fc)
    r2=par.add_run(); it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text=instr; r2._r.append(it)
    r3=par.add_run(); fc2=OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'),'separate'); r3._r.append(fc2)
    r4=par.add_run("(alanı güncellemek için F9)");
    r5=par.add_run(); fc3=OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'),'end'); r5._r.append(fc3)

def render_table(doc, rows):
    cells=[[c.strip() for c in r.strip().strip('|').split('|')] for r in rows]
    cells=[r for i,r in enumerate(cells) if not (set(rows[i].strip()) <= set('|-: ') and '-' in rows[i])]
    if not cells: return
    ncol=max(len(r) for r in cells)
    t=doc.add_table(rows=0, cols=ncol); t.style='Light Grid Accent 1'
    for ri,row in enumerate(cells):
        row=row+['']*(ncol-len(row)); tr=t.add_row().cells
        for ci,val in enumerate(row):
            p=tr[ci].paragraphs[0]; add_runs(p, val)
            for rn in p.runs: rn.font.size=Pt(8.5)
            if ri==0:
                set_cell_bg(tr[ci],'D9E2F3')
                for rn in p.runs: rn.bold=True

def main():
    doc=Document()
    st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(11)
    st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.15
    # dil tr-TR
    rpr=st.element.get_or_add_rPr(); lang=OxmlElement('w:lang'); lang.set(qn('w:val'),'tr-TR'); rpr.append(lang)
    for s in ['Heading 1','Heading 2','Heading 3']:
        try: doc.styles[s].font.color.rgb=RGBColor(0x1F,0x38,0x64)
        except: pass
    sec=doc.sections[0]; sec.left_margin=sec.right_margin=Inches(0.9)
    # altbilgi sayfa no
    ftr=sec.footer.paragraphs[0]; ftr.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=ftr.add_run(); f=OxmlElement('w:fldSimple'); f.set(qn('w:instr'),'PAGE'); r._r.addprevious(f)
    ftr.add_run(" / "); r2=ftr.add_run(); f2=OxmlElement('w:fldSimple'); f2.set(qn('w:instr'),'NUMPAGES'); r2._r.addprevious(f2)
    # üstbilgi
    sec.header.paragraphs[0].text="KVKK Belge Paketi — TASLAK (hukukçu onayı bekliyor)"
    for rn in sec.header.paragraphs[0].runs: rn.font.size=Pt(8); rn.italic=True

    # KAPAK
    for _ in range(6): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("KVKK BELGE PAKETİ"); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor(0x1F,0x38,0x64)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Mentor–Menti Eşleştirme Platformu"); r.font.size=Pt(14)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Sürüm tarihi: 25.08.2026"); r.font.size=Pt(12)
    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("TASLAK — HUKUKÇU ONAYI BEKLİYOR"); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=RGBColor(0xC0,0x00,0x00)
    doc.add_page_break()

    # NASIL İNCELENSİN
    doc.add_heading("Bu Paket Nasıl İncelenmeli?",level=1)
    for line in [
        "Bu paket, platformun kişisel veri işleme faaliyetlerine ilişkin yasal belge taslaklarını içerir. Metinler, sistemin FİİLEN yaptığı işlemlere (kod gerçeğine) dayanılarak sıfırdan hazırlanmıştır; jenerik şablon değildir.",
        "• Lütfen belgeyi 'Değişiklikleri İzle' (Track Changes) açık şekilde düzenleyiniz.",
        "• SARI zeminli alanlar sizin hukuki kararınızı bekleyen noktalardır ([HUKUKÇU KARARI: ...]).",
        "• GRİ zeminli alanları ([... DOLDURACAK: ...]) veri sorumlusu/işletmeci dolduracaktır.",
        "• Onayınız ve düzenlemeleriniz sonrası bu işaretler kaldırılıp metinler yayınlanacaktır.",
        "• İçindekiler tablosunu güncellemek için tabloya sağ tıklayıp 'Alanı Güncelle' (F9) seçiniz.",
    ]:
        pp=doc.add_paragraph(); add_runs(pp,line)
    doc.add_page_break()

    # İÇİNDEKİLER
    doc.add_heading("İçindekiler",level=1)
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()

    # BELGELER
    for idx,fn in enumerate(ORDER):
        path=os.path.join(BASE,fn)
        with open(path,encoding="utf-8") as fh: lines=fh.read().split("\n")
        doc.add_heading(deemoji(TITLES[fn]),level=1)
        i=0
        while i<len(lines):
            ln=lines[i]
            if ln.strip().startswith("|") and i+1<len(lines) and (set(lines[i+1].strip())<=set('|-: ') and '-' in lines[i+1]):
                block=[]
                while i<len(lines) and lines[i].strip().startswith("|"): block.append(lines[i]); i+=1
                render_table(doc,block); continue
            s=ln.strip()
            if not s: i+=1; continue
            if s.startswith("#### "): doc.add_heading(deemoji(s[5:]),level=4)
            elif s.startswith("### "): doc.add_heading(deemoji(s[4:]),level=3)
            elif s.startswith("## "): doc.add_heading(deemoji(s[3:]),level=2)
            elif s.startswith("# "): doc.add_heading(deemoji(s[2:]),level=2)
            elif s.startswith(">"):
                pp=doc.add_paragraph(); add_runs(pp, s.lstrip('> ').rstrip())
                pp.paragraph_format.left_indent=Inches(0.2)
                for rn in pp.runs: rn.font.size=Pt(9.5); rn.italic=True
            elif re.match(r'^[-*] ',s):
                pp=doc.add_paragraph(style='List Bullet'); add_runs(pp,s[2:])
            elif re.match(r'^\d+\. ',s):
                pp=doc.add_paragraph(style='List Number'); add_runs(pp,re.sub(r'^\d+\. ','',s))
            elif set(s)<=set('-'): pass
            else:
                pp=doc.add_paragraph(); add_runs(pp,s)
            i+=1
        if idx<len(ORDER)-1: doc.add_page_break()

    doc.save(OUT)
    print("YAZILDI:", OUT)

if __name__=="__main__": main()
